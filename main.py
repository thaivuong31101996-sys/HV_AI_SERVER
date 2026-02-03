from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse
import docx
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, io, re, logging

# 1. CẤU HÌNH NHẬT KÝ ĐỂ THEO DÕI LỖI TRÊN RENDER
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

def has_image(p):
    """Kiểm tra xem paragraph có chứa ảnh không"""
    return True if p._element.xpath('.//w:drawing') or p._element.xpath('.//w:inline') else False

@app.get("/")
async def home():
    return {"message": "Shop Hai Vuong AI is Ready!"}

@app.post("/process")
async def process_word(
    file: UploadFile = File(...),
    le_trai: float = Form(3.0), le_phai: float = Form(2.0),
    le_tren: float = Form(2.0), le_duoi: float = Form(2.0),
    h1_size: int = Form(14), h2_size: int = Form(13), h3_size: int = Form(13),
    line_spacing: float = Form(1.3), indent: float = Form(1.27),
    h1_upper: bool = Form(True), chua_trang: int = Form(1)
):
    try:
        logger.info(f"--- Đang xử lý file: {file.filename} ---")
        content = await file.read()
        old_doc = Document(io.BytesIO(content))
        
        # TẠO FILE MỚI SẠCH 100%
        new_doc = Document()
        
        # Ép lề chuẩn shop Hải Vương
        section = new_doc.sections[0]
        section.left_margin, section.right_margin = Cm(le_trai), Cm(le_phai)
        section.top_margin, section.bottom_margin = Cm(le_tren), Cm(le_duoi)

        for p in old_doc.paragraphs:
            txt = p.text.strip()
            # Bỏ dòng trống không cần thiết
            if not txt and not has_image(p): continue 

            new_p = new_doc.add_paragraph()
            
            # --- AI PHÂN TÍCH VÀ ĐỘ FILE ---
            
            # Cấp 1: Chương và các mục tương đương
            if re.match(r'^(CHƯƠNG|LỜI MỞ ĐẦU|KẾT LUẬN|DANH MỤC|PHỤ LỤC|TÀI LIỆU|MỤC LỤC|PHẦN)', txt.upper()):
                new_p.text = txt.upper() if h1_upper else txt
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_p.paragraph_format.space_before = Pt(18)
                cur_size, is_bold = h1_size, True
            
            # Cấp 2: Mục 1.1, 1.2 (Sát lề gáy)
            elif re.match(r'^\d+\.\d+(\s|$)', txt):
                new_p.text = txt
                new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_p.paragraph_format.space_before = Pt(12)
                cur_size, is_bold = h2_size, True
                
            # Cấp 3: Mục 1.2.1 (Thụt vào đúng bằng Thụt đầu dòng nội dung)
            elif re.match(r'^\d+\.\d+\.\d+', txt):
                new_p.text = txt
                new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_p.paragraph_format.left_indent = Cm(indent)
                cur_size, is_bold = h3_size, True

            # Ảnh và chú thích
            elif has_image(p) or re.match(r'^(Bảng|Hình|Ảnh|Table|Figure)', txt, re.IGNORECASE):
                new_p.text = txt
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cur_size, is_bold = 12, False

            # Văn bản thường (Body)
            else:
                new_p.text = txt
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_p.paragraph_format.line_spacing = line_spacing
                new_p.paragraph_format.first_line_indent = Cm(indent)
                cur_size, is_bold = 13, False

            # GÁN FONT TIMES NEW ROMAN TRIỆT ĐỂ
            if new_p.runs:
                for run in new_p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(cur_size)
                    run.bold = is_bold
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            else:
                run = new_p.add_run(txt)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(cur_size)
                run.bold = is_bold
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # XUẤT FILE VỀ MÁY TẠI TIỆM
        out_io = io.BytesIO()
        new_doc.save(out_io)
        out_io.seek(0)
        logger.info("Xử lý thành công, đang gửi file về!")
        return FileResponse(out_io, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="HV_DONE.docx")

    except Exception as e:
        logger.error(f"LỖI TẠI SERVER: {str(e)}")
        return {"error": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=10000)
